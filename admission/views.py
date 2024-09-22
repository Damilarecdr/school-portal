from django.shortcuts import render, redirect, get_object_or_404
from .models import Application, School_Fee_Payment, Session
from .forms import ApplicationForm, ClassUpgradeForm
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
import csv
from django.http import HttpResponse
from django.template.loader import render_to_string
from datetime import datetime
from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.http import HttpResponse
from xhtml2pdf import pisa
from .models import Application
from django.http import HttpResponse
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from .models import Application
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from django.contrib.admin.views.decorators import staff_member_required
from django.utils import timezone
from datetime import datetime
from .models import Application
from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required

from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from datetime import datetime

from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Course, Registration
import csv

import requests
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from .models import Application, School_Fee_Payment

from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from .models import Application, School_Fee_Payment
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from django.shortcuts import render
from .models import Registration

from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import Application, Session  # Adjust the import based on your project structure

@login_required
def dashboard(request):
    application = Application.objects.filter(user=request.user).first()
    current_session = Session.objects.first()  # Adjust this if you have a specific way to get the current session

    if application and application.is_approved:
        # List of classes that have special dashboards
        special_classes = [
            'CREACH', 'KG 1', 'KG 2', 'NURSERY 1', 'NURSERY 2', 'NURSERY 3',
            'PRIMARY 1', 'PRIMARY 2', 'PRIMARY 3', 'PRIMARY 4', 'PRIMARY 5'
        ]

        # Check if the application class is in the special classes
        if application.class_applying in special_classes:
            return render(request, 'admission/primary_dashboard.html', {
                'application': application,
                'current_session': current_session
            })
        else:
            # Render a different dashboard for other approved applications
            return render(request, 'student_dashboard.html', {
                'application': application,
                'current_session': current_session
            })

    # For applications that are not approved or if there's no application
    return render(request, 'admission/app_dashboard.html', {'application': application})



@login_required
def apply(request):
    if request.method == 'POST':
        form = ApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            application = form.save(commit=False)
            application.user = request.user
            application.save()
            return redirect('dashboard')
    else:
        form = ApplicationForm()
    return render(request, 'admission/apply.html', {'form': form})

@staff_member_required
def admin_view_applications(request):
    if 'download_csv' in request.GET:
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="applications_by_class.csv"'
        writer = csv.writer(response)

        class_groups = ['JSS 1', 'JSS 2', 'JSS 3', 'SSS 1', 'SSS 2', 'SSS 3']
        writer.writerow(['Class', 'Surname', 'Other Names', 'Email', 'Approved'])

        for class_group in class_groups:
            writer.writerow([f'Class: {class_group}'])
            applications = Application.objects.filter(class_applying=class_group)

            for application in applications:
                writer.writerow([
                    '',
                    application.surname,
                    application.other_names,
                    application.email,
                    'Yes' if application.is_approved else 'No'
                ])

        return response

    if request.method == 'POST' and 'csv_upload' in request.FILES:
        csv_file = request.FILES['csv_upload']
        if not csv_file.name.endswith('.csv'):
            messages.error(request, 'This is not a valid CSV file.')
            return redirect('admin_view_applications')

        decoded_file = csv_file.read().decode('utf-8').splitlines()
        reader = csv.DictReader(decoded_file)

        for row in reader:
            try:
                email = row['Email']
                application = Application.objects.get(user__email=email)
                application.is_approved = row['Approved'].strip().lower() == 'yes'
                application.save()
            except Application.DoesNotExist:
                messages.error(request, f"No application found for email {email}")
                continue
        messages.success(request, 'Applications updated successfully.')

    applications = Application.objects.filter(is_approved=False)
    return render(request, 'admission/admin_applications.html', {'applications': applications})




from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import get_object_or_404, redirect
from django.utils import timezone
from datetime import datetime
from .models import Application  # Adjust import based on your structure

'''@staff_member_required
def approve_application(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    application.is_approved = True

    # Generate admission number based on class applying
    current_year = datetime.now().year

    # Define the classes for the special admission number format
    special_classes = [
        'KG 1', 'KG 2', 'NURSERY 1', 'NURSERY 2', 'NURSERY 3',
        'PRIMARY 1', 'PRIMARY 2', 'PRIMARY 3', 'PRIMARY 4', 'PRIMARY 5'
    ]

    if application.class_applying in special_classes:
        application.admission_number = f"GHSPE/{current_year}/{application.id:04d}"
    else:
        application.admission_number = f"GHSSE/{current_year}/{application.id:04d}"

    # Set the approval date
    application.date_admission_approved = timezone.now()

    # Change the username to the admission number
    user = application.user
    user.username = application.admission_number
    user.save()

    application.save()

    return redirect('applications')'''

@staff_member_required
def approve_application(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    application.is_approved = True

    # Set the current year
    current_year = datetime.now().year

    # Define classes for special admission number format
    nursery_primary_classes = [
        'KG 1', 'KG 2', 'NURSERY 1', 'NURSERY 2', 'NURSERY 3',
        'PRIMARY 1', 'PRIMARY 2', 'PRIMARY 3', 'PRIMARY 4', 'PRIMARY 5'
    ]
    
    secondary_classes = ['JSS 1', 'JSS 2', 'JSS 3', 'SSS 1', 'SSS 2', 'SSS 3']

    # Determine the base number format based on class applying
    if application.class_applying in nursery_primary_classes:
        base_number = f"GHSPE/{current_year}/"
        count = Application.objects.filter(
            class_applying__in=nursery_primary_classes,
            is_approved=True
        ).count() + 1  # Start counting from the current number of approved applications

        # Generate admission number
        admission_number = f"{base_number}{count:04d}"
        while Application.objects.filter(admission_number=admission_number).exists():
            count += 1
            admission_number = f"{base_number}{count:04d}"

        application.admission_number = admission_number
        
    elif application.class_applying in secondary_classes:
        base_number = f"GHSSE/{current_year}/"
        count = Application.objects.filter(
            class_applying__in=secondary_classes,
            is_approved=True
        ).count() + 1  # Start counting from the current number of approved applications

        # Generate admission number
        admission_number = f"{base_number}{count:04d}"
        while Application.objects.filter(admission_number=admission_number).exists():
            count += 1
            admission_number = f"{base_number}{count:04d}"

        application.admission_number = admission_number

    # Set the approval date
    application.date_admission_approved = timezone.now()

    # Change the username to the admission number
    user = application.user
    user.username = application.admission_number
    user.save()

    application.save()

    return redirect('applications')


'''@staff_member_required
def approve_application(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    application.is_approved = True

    # Set the current year
    current_year = datetime.now().year

    # Define classes for special admission number format
    nursery_primary_classes = [
        'KG 1', 'KG 2', 'NURSERY 1', 'NURSERY 2', 'NURSERY 3',
        'PRIMARY 1', 'PRIMARY 2', 'PRIMARY 3', 'PRIMARY 4', 'PRIMARY 5'
    ]

    # Generate admission number based on class applying
    if application.class_applying in nursery_primary_classes:
        # Count existing admissions for nursery/primary classes for unique numbering
        count = Application.objects.filter(
            class_applying__in=nursery_primary_classes,
            is_approved=True
        ).count() + 1  # Add 1 to include the current application

        application.admission_number = f"GHSPE/{current_year}/{application.id:01d}/{count:04d}"
    else:
        # Count existing admissions for secondary classes for unique numbering
        count = Application.objects.filter(
            class_applying__in=['JSS 1', 'JSS 2', 'JSS 3', 'SSS 1', 'SSS 2', 'SSS 3'],
            is_approved=True
        ).count() + 1  # Add 1 to include the current application

        application.admission_number = f"GHSSE/{current_year}/{application.id:01d}/{count:04d}"

    # Set the approval date
    application.date_admission_approved = timezone.now()

    # Change the username to the admission number
    user = application.user
    user.username = application.admission_number
    user.save()

    application.save()

    return redirect('applications')'''


@login_required
def print_application(request):
    application = Application.objects.filter(user=request.user).first()
    current_session = Session.objects.first()
    if not application:
        return HttpResponse("No application found.")

    context = {
        'application': application,
        'current_session': current_session,
        'admission_number': application.admission_number,
        'current_year': datetime.now().year,
    }

    # Render HTML template with application details
    html = render_to_string('application_pdf.html', context)
    
    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="application_{application.admission_number}.pdf"'
    
    # Generate PDF from HTML
    pisa_status = pisa.CreatePDF(html, dest=response)

    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    
    return response


@login_required
def admission_letter(request):
    application = Application.objects.filter(user=request.user).first()
    payment = School_Fee_Payment.objects.filter(application=application).first()
    current_session = Session.objects.first()

    current_year = datetime.now().year
    admission_number = f"GHSSE/{current_year}/{application.id:04d}"
    if not application:
        return HttpResponse("No application found.")

    context = {
        'application': application,
        'admission_number': admission_number,
        'current_year': current_year,
    }
    html = render_to_string('admission/ad.html', context)
    # Render HTML template with application details
    
    
    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="application_{application.admission_number}.pdf"'
    
    # Generate PDF from HTML
    pisa_status = pisa.CreatePDF(html, dest=response)

    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')
    
    return response



'''@login_required
def admission_letter(request):
    application = Application.objects.filter(user=request.user).first()
    payment = School_Fee_Payment.objects.filter(application=application).first()

    #if not payment or not payment.is_paid:
     #   return HttpResponse("You must pay the school fee before printing your admission letter.")

    current_year = datetime.now().year
    admission_number = f"GHSSE/{current_year}/{application.id:04d}"

    context = {
        'application': application,
        'admission_number': admission_number,
        'current_year': current_year,
    }
    html = render_to_string('admission/ad.html', context)
    
    return HttpResponse(html)'''

@staff_member_required
def download_approved_admissions(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="approved_admissions.csv"'

    writer = csv.writer(response)
    writer.writerow(['Surname', 'Other Names', 'Email', 'Class Applying', 'State of Origin', 'LGA', 'Parent Name', 'Parent Email'])

    approved_applications = Application.objects.filter(is_approved=True)

    for application in approved_applications:
        writer.writerow([
            application.surname,
            application.other_names,
            application.email,
            application.class_applying,
            application.state_of_origin,
            application.lga,
            application.parent_name,
            application.parent_email,
        ])

    return response

@login_required
def general_payment(request):
    return render(request, 'payment/general_payments.html')


'''@login_required
def payment(request):
    application = get_object_or_404(Application, user=request.user)
    payment, created = School_Fee_Payment.objects.get_or_create(application=application)

    # Determine the school fee based on the class applying to
    if "JSS" in application.class_applying:
        school_fee = 6000
    else:
        school_fee = 10000

    if request.method == 'POST':
        # Mark the payment as paid
        payment.is_paid = True
        payment.school_fee = school_fee  # Set the school fee from the application
        payment.save()
        messages.success(request, 'Payment marked as paid successfully.')
        return redirect('dashboard')

    context = {
        'application': application,
        'payment': payment,
        'school_fee': school_fee,
    }
    return render(request, 'payment/payment.html', context)'''


@login_required
def payment(request):
    application = get_object_or_404(Application, user=request.user)
    payment, created = School_Fee_Payment.objects.get_or_create(application=application)

    # Determine the school fee and breakdown based on the class applying to
    if "JSS" in application.class_applying:
        school_fee = 32500
        breakdown = {
            "Form": 1500,
            "School Uniform": 6000,
            "School Fee": 13000,
            "Exam Fees": 2000,
            "Lesson Fees": 2000,
            "Card": 1500,
            "Sport Wears": 6000,
            "Portal Charges": 500,
            "Total": 32500,
        }
    elif "SSS" in application.class_applying:
        school_fee = 36500
        breakdown = {
            "Form": 1500,
            "School Uniform": 7000,
            "School Fee": 16000,
            "Exam Fees": 2000,
            "Lesson Fees": 2000,
            "Card": 1500,
            "Sport Wears": 6000,
            "Portal Charges": 500,
            "Total": 36500,
        }
    elif "PRIMARY" in application.class_applying:
        school_fee = 23200  # Example fee for primary
        breakdown = {
            "Form": 1000,
            "School Uniform": 4500,
            "School Fee": 8000,
            "Exam Fee": 1000,
            "Lesson Fees": 1000,
            "Card": 1500,
            "Sport Wears": 6000,
            "Portal Charges": 200,
            "Total": 23200,
        }
    elif "NURSERY" in application.class_applying:
        school_fee = 19150  # Example fee for nursery
        breakdown = {
            "Form": 1000,
            "School Uniform": 4000,
            "School Fee": 7000,
            "Exam Fee": 1000,
            "Lesson Fees": 1000,
            "Card": 1000,
            "Sport Wears": 4000,
            "Portal Charges": 150,
            "Total": 19150,
        }
    elif "KG" in application.class_applying:
        school_fee = 17100  # Example fee for nursery
        breakdown = {
            "Form": 1000,
            "School Uniform": 3500,
            "School Fee": 6000,
            "Exam Fee": 1000,
            "Lesson Fees": 1000,
            "Card": 1000,
            "Sport Wears": 4000,
            "Portal Charges": 100,
            "Total": 17100,
        }
    elif "CREACH" in application.class_applying:
        school_fee = 16100  # Example fee for nursery
        breakdown = {
            "Form": 1000,
            "School Uniform": 3500,
            "School Fee": 6000,
            "Exam Fee": 1000,
            
            "Card": 1000,
            "Sport Wears": 4000,
            "Portal Charges": 100,
            "Total": 16100,
        }
    else:
        # Handle other cases or raise an error
        school_fee = 10000
        breakdown = {
            "Tuition Fee": 7000,
            "Activity Fee": 3000,
        }

    if request.method == 'POST':
        # Initiate a payment with Paystack
        payment_data = {
            'email': request.user.email,
            'amount': school_fee * 100,  # Paystack accepts amount in kobo
            'callback_url': 'http://127.0.0.1:8000/payment/callback/',  # Update with your callback URL
        }

        response = requests.post(
            'https://api.paystack.co/transaction/initialize',
            json=payment_data,
            headers={'Authorization': f'Bearer {settings.PAYSTACK_SECRET_KEY}'}
        )

        response_data = response.json()
        if response_data['status']:
            return redirect(response_data['data']['authorization_url'])
        else:
            messages.error(request, 'Payment initiation failed.')

    context = {
        'application': application,
        'payment': payment,
        'school_fee': school_fee,
        'breakdown': breakdown,
    }
    return render(request, 'payment/payment.html', context)


@csrf_exempt
def payment_callback(request):
    if request.method == 'GET':
        payment_reference = request.GET.get('reference')
        
        # Verify the payment
        response = requests.get(
            f'https://api.paystack.co/transaction/verify/{payment_reference}',
            headers={'Authorization': f'Bearer {settings.PAYSTACK_SECRET_KEY}'}
        )
        response_data = response.json()
        
        if response_data['status'] and response_data['data']['status'] == 'success':
            # Payment successful
            application = get_object_or_404(Application, user=request.user)
            payment, created = School_Fee_Payment.objects.get_or_create(application=application)
            payment.is_paid = True
            payment.school_fee = response_data['data']['amount'] / 100  # Convert kobo to naira
            payment.save()

            messages.success(request, 'Payment successful!')
            return redirect('dashboard')
        else:
            messages.error(request, 'Payment verification failed.')

    return redirect('dashboard')





@login_required
def payment_receipt(request):
    # Retrieve the application for the logged-in user
    application = get_object_or_404(Application, user=request.user)

    # Retrieve the corresponding school fee payment
    school_fee_payment = get_object_or_404(School_Fee_Payment, application=application)

    # Check if the payment has been made
    if not payment or not school_fee_payment.is_paid:
        return redirect('payment')

    # Define fee breakdown based on class
    if "JSS" in application.class_applying:
        breakdown = {
            "Form": 1500,
            "School Uniform": 6000,
            "School Fee": 13000,
            "Exam Fees": 2000,
            "Lesson Fees": 2000,
            "Card": 1500,
            "Sport Wears": 6000,
            "Portal Charges": 500,
            "Total": 32500,
        }
    elif "SSS" in application.class_applying:
        breakdown = {
            "Form": 1500,
            "School Uniform": 7000,
            "School Fee": 16000,
            "Exam Fees": 2000,
            "Lesson Fees": 2000,
            "Card": 1500,
            "Sport Wears": 6000,
            "Portal Charges": 500,
            "Total": 36500,
        }
    else:
        breakdown = {
            "Tuition Fee": 7000,
            "Activity Fee": 3000,
            "Total": 10000,
        }

    # Calculate totals
    net_total = breakdown["Total"]
 

    # Prepare context for rendering the template
    context = {
        'application': application,
        'school_fee_payment': school_fee_payment,
        'breakdown': breakdown,
        'net_total': net_total,
      
    }

    
    # Render the HTML for the receipt
    return render(request, 'payment/payment_receipt.html', context)

    #html = render_to_string('payment/payment_receipt.html', context)

    '''# Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="payment_receipt_{application.admission_number}.pdf"'
    
    # Generate PDF from HTML
    pisa_status = pisa.CreatePDF(html, dest=response)

    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html + '</pre>')

    return response'''









@login_required
def list_courses(request):
    courses = Course.objects.all()
    return render(request, 'list_course.html', {'courses': courses})

@login_required
def register_courses(request):
    application = Application.objects.filter(user=request.user).first()
    payment = School_Fee_Payment.objects.filter(application=application).first()
    if request.method == "POST":
        selected_courses = request.POST.getlist('courses')  # Get selected course IDs
        for course_id in selected_courses:
            course = Course.objects.get(id=course_id)
            Registration.objects.get_or_create(user=request.user, course=course)
    if not payment or not payment.is_paid:
        return HttpResponse("You must pay the school fee before printing your admission letter.")
        return redirect('registration_success')  # Redirect to a success page
    return redirect('list_courses')  # Redirect if not POST
    from django.shortcuts import render

@login_required
def registration_success(request):
    return render(request, 'courseregsuc.html')

@staff_member_required
def download_registrations(request, course_id):
    course = Course.objects.get(id=course_id)
    registrations = Registration.objects.filter(course=course)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{course.name}_registrations.csv"'

    writer = csv.writer(response)
    writer.writerow(['Username', 'Email', 'Registration Date'])
    
    for registration in registrations:
        writer.writerow([registration.user.username, registration.user.email, registration.registration_date])

    return response


@login_required
def registered_courses(request):
    if request.user.is_authenticated:
        registrations = Registration.objects.filter(user=request.user)
        return render(request, 'printcourse.html', {'registrations': registrations})
    else:
        return redirect('list_courses')  # Redirect to courses if user is not authenticated


from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from io import BytesIO
from django.templatetags.static import static
import os
from django.conf import settings

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.admin.views.decorators import staff_member_required
from django.core.files.storage import FileSystemStorage
from docx import Document
import os
from io import BytesIO
from .models import Application, User
from django.conf import settings


@staff_member_required
def view_approved_applications(request):
    approved_applications = Application.objects.filter(is_approved=True).order_by('class_applying')

    # Check if the user requested a DOCX download
    if 'download' in request.GET:
        # Create a new Document
        document = Document()

        # Add the main heading: "God's Heritage School"
        document.add_heading("God's Heritage School", level=0)

        # Group applications by class
        class_groups = {}
        for application in approved_applications:
            if application.class_applying not in class_groups:
                class_groups[application.class_applying] = []
            class_groups[application.class_applying].append(application)

        # Add a heading and details for each class
        for class_applying, applications in class_groups.items():
            # Add a heading for each class
            document.add_heading(f"Class: {class_applying}", level=1)

            # Add a table for each class with the necessary columns
            table = document.add_table(rows=1, cols=8)
            table.style = 'Table Grid'  # Optional: Apply grid style to make it look tabular
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Surname'
            hdr_cells[1].text = 'Other Names'
            hdr_cells[2].text = 'State of Origin'
            hdr_cells[3].text = 'Local Government Area'
            hdr_cells[4].text = 'Parent Email'
            hdr_cells[5].text = 'Parent Phone Number'
            hdr_cells[6].text = 'Admission Number'
            hdr_cells[7].text = 'Class'
            
           

            # Add data rows
            for application in applications:
                row_cells = table.add_row().cells
                row_cells[0].text = application.surname
                row_cells[1].text = application.other_names
                row_cells[2].text = application.state_of_origin
                row_cells[3].text = application.lga
                row_cells[4].text = application.parent_email
                row_cells[5].text = application.parent_phone
                row_cells[6].text = application.user.username 
                row_cells[7].text = application.class_applying

                for cell in row_cells:
                    cell.paragraphs[0].space_after = 100  # Space after the paragraph
                    cell.paragraphs[0].space_before = 100  # Space before the paragraph
                    cell.paragraphs[0].paragraph_format.alignment = 1  # Center-align tex
            # Optional: Set the width for each column for better spacing
        for col in table.columns:
             for cell in col.cells:
                cell.width = Inches(1.5)  # Adjust width as needed
                

        # Prepare the response as a docx file
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="approved_applications.docx"'

        # Save the document to a BytesIO object
        doc_io = BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        # Write the document to the response
        response.write(doc_io.read())
        return response

    # If there's a file upload
    if request.method == 'POST' and request.FILES['upload']:
        upload_file = request.FILES['upload']
        fs = FileSystemStorage()
        filename = fs.save(upload_file.name, upload_file)
        uploaded_file_path = fs.path(filename)

        # Read and process the uploaded DOCX file
        document = Document(uploaded_file_path)
        table = document.tables[0]  # Assuming there's one table in the document

        # Iterate over the rows (skipping the first header row)
        for row in table.rows[1:]:
            surname = row.cells[0].text.strip()
            other_names = row.cells[1].text.strip()
            email = row.cells[2].text.strip()
            class_applying = row.cells[3].text.strip()
            username = row.cells[4].text.strip()

            # Try to find the user and application based on the username or email
            try:
                user = User.objects.get(username=username)
                application = Application.objects.get(user=user)

                # Update the user and application details
                user.email = email
                user.save()

                application.surname = surname
                application.other_names = other_names
                application.class_applying = class_applying
                application.save()
            except User.DoesNotExist:
                # If the user is not found, skip or log an error
                continue

        # Redirect to refresh the page after upload
        return redirect('students')

    return render(request, 'admission/view_student.html', {'approved_applications': approved_applications})



@staff_member_required
def change_student_class(request, application_id):
    application = get_object_or_404(Application, id=application_id)
    
    if request.method == 'POST':
        new_class = request.POST.get('class_applying')
        if new_class:
            application.class_applying = new_class
            application.save()
            messages.success(request, 'Class updated successfully.')
        else:
            messages.error(request, 'Please select a class.')

    return redirect('students')

